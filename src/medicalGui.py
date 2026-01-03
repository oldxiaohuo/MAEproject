import os
import sys
import pydicom
import numpy as np
from PyQt5.QtWidgets import (QApplication, QMainWindow, QPushButton, QFileDialog, QLabel,QHBoxLayout,
                             QComboBox, QTextEdit, QListWidget, QVBoxLayout, QWidget, QProgressBar,
                             QMessageBox)
from PyQt5.QtGui import QImage, QPixmap, QIcon
from PyQt5.QtCore import Qt
from docx import Document

os.environ["QT_QPA_PLATFORM_PLUGIN_PATH"] = r'H:\Users\小巷鱼\AppData\Local\Programs\Python\Python313\Lib\site-packages\PyQt5\Qt5\plugins\platforms'

class DummyModel:
    def __init__(self):
        self.models = {
            "模型A": self.predict_a,
            "模型B": self.predict_b,
            "模型C": self.predict_c
        }

    def predict(self, model_name, image):
        return self.models[model_name](image)

    def predict_a(self, image):
        return image * 0.8

    def predict_b(self, image):
        return image * 0.5

    def predict_c(self, image):
        return image * 1.2


class MedicalGUI(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle('医学AI报告软件')
        self.setGeometry(100, 100, 1000, 700)

        icon = QIcon()
        icon.addPixmap(QPixmap("./favicon.png"), QIcon.Normal, QIcon.Off)
        self.setWindowIcon(icon)  # 替换成你的 ico 路径
    

        self.model = DummyModel()
        self.patient_info = {}

        self.file_list = QListWidget()
        self.file_list.itemClicked.connect(self.display_dicom)

        self.model_selector = QComboBox()
        self.model_selector.addItems(self.model.models.keys())

        self.info_label = QLabel('患者信息：')
        self.progress_bar = QProgressBar()

        # 主布局
        main_layout = QVBoxLayout()

        # 顶部横向布局
        top_layout = QHBoxLayout()

        # 左侧图像
        self.image_label = QLabel()
        self.image_label.setFixedSize(256, 512)
        self.image_label.setStyleSheet('border: 1px solid black')

        # 右侧报告
        self.report_edit = QTextEdit()

        # 添加到横向布局
        top_layout.addWidget(self.image_label)
        top_layout.addWidget(self.report_edit)

        # 其他控件
        main_layout.addWidget(self.model_selector)
        main_layout.addWidget(self.file_list)
        main_layout.addWidget(self.info_label)
        main_layout.addLayout(top_layout)  # 把横向布局放进主布局
        main_layout.addWidget(self.progress_bar)

        container = QWidget()
        container.setLayout(main_layout)
        self.setCentralWidget(container)


        # 添加菜单栏
        self.setup_menu()

        self.dicom_files = []
        self.current_prediction = None

    def setup_menu(self):
        menu_bar = self.menuBar()

        file_menu = menu_bar.addMenu('文件')
        file_menu.addAction('打开 DICOM 文件夹', self.load_folder)
        file_menu.addAction('加载报告模板', self.load_template)
        file_menu.addAction('退出程序', self.close)

        edit_menu = menu_bar.addMenu('编辑')
        edit_menu.addAction('清空报告', self.clear_report)
        edit_menu.addAction('刷新进度', self.reset_progress)

        tools_menu = menu_bar.addMenu('工具')
        tools_menu.addAction('生成报告', lambda: self.generate_report(self.current_prediction))
        tools_menu.addAction('显示关于', self.show_about)

    def load_folder(self):
        folder_path = QFileDialog.getExistingDirectory(self, '选择DICOM文件夹')
        if folder_path:
            self.dicom_files = [os.path.join(folder_path, f) for f in os.listdir(folder_path) if (f.endswith('.dcm') or f.endswith('.DCM'))]
            self.file_list.clear()
            self.file_list.addItems(self.dicom_files)

    def load_template(self):
        template_path, _ = QFileDialog.getOpenFileName(self, '选择Word模板', '', 'Word Files (*.docx)')

        replacements = {
            '<姓名>': self.patient_info.get('姓名', ''),
            '<检查ID>': self.patient_info.get('检查ID', ''),
            '<性别>': self.patient_info.get('性别', ''),
            '<出生日期>': self.patient_info.get('出生日期', ''),
            '<检查项目>': self.patient_info.get('检查项目', ''),
            '<检查日期>': self.patient_info.get('检查日期', ''),
            # '<分析结果>': str(prediction_result)
        }

        if template_path:
            self.template_doc = Document(template_path)
            text = '\n'.join([para.text for para in self.template_doc.paragraphs])
        
            for key, value in replacements.items():
                if key in text:
                    text = text.replace(key, value)
            text = text + f'\n分析结果：预测病灶面积 {self.current_prediction:.2f}'

            self.report_edit.setPlainText(text)
            self.template_path = template_path
            QMessageBox.information(self, '模板加载成功', '模板已成功加载，请继续编辑或生成报告。')
    
    
    def display_dicom(self, item):
        dicom_path = item.text()
        ds = pydicom.dcmread(dicom_path)
        image = ds.pixel_array.astype(np.float32)
        
        if ds.Modality == 'NM':

            anterior = image[0]
            posterior = image[1]
            anterior = np.squeeze(anterior)
            posterior = np.squeeze(posterior)
            image = np.hstack((anterior,posterior))


        patient_id = ds.PatientID
        patient_name = ds.PatientName
        patient_sex = ds.PatientSex
        patient_birthday = ds.PatientBirthDate
        study_date = getattr(ds, 'StudyDate', '未知')
        study_id = getattr(ds, 'StudyID', '未知')

        # charset = getattr(ds, 'SpecificCharacterSet', 'utf-8')
        # print(charset)
        # def safe_read(field, charset):
        #     if isinstance(field, bytes):
        #         try:
        #             return field.decode(charset)
        #         except:
        #             return field.decode('utf-8', errors='ignore')
        #     return str(field)
        
        # patient_sex = safe_read(ds.PatientSex, charset)
        # patient_birthday = safe_read(ds.PatientBirthDate, charset)

        self.patient_info = {
            '检查ID': patient_id,
            '姓名': str(patient_name),
            '性别': patient_sex,
            '出生日期': patient_birthday,
            '检查项目': study_id,
            '检查日期': study_date
        }

        self.info_label.setText(f'检查ID: {patient_id} | 姓名: {patient_name} | 性别: {patient_sex} | 出生日期: {patient_birthday} | '
                                f'检查项目: {study_id} | 检查日期: {study_date}')

        model_name = self.model_selector.currentText()
        prediction = self.model.predict(model_name, image)

        lesion_area = np.sum(prediction > 0.5)
        self.current_prediction = lesion_area

        self.show_image_on_label(prediction)

        self.progress_bar.setValue(100)
        self.report_edit.append(f'\n分析结果：预测病灶面积 {lesion_area:.2f}')

    def show_image_on_label(self, image_array):
        norm_img = ((image_array - np.min(image_array)) / (np.max(image_array) - np.min(image_array)) * 255).astype(np.uint8)
        h, w = norm_img.shape
        q_image = QImage(norm_img.data, w, h, w, QImage.Format_Grayscale8)
        pixmap = QPixmap.fromImage(q_image)
        self.image_label.setPixmap(pixmap.scaled(self.image_label.size(), Qt.KeepAspectRatio))

    def generate_report(self, prediction_result):
        if not hasattr(self, 'template_path'):
            QMessageBox.warning(self, '错误', '请先加载报告模板。')
            return

        replacements = {
            '<姓名>': self.patient_info.get('姓名', ''),
            '<检查ID>': self.patient_info.get('检查ID', ''),
            '<性别>': self.patient_info.get('性别', ''),
            '<出生日期>': self.patient_info.get('出生日期', ''),
            '<检查项目>': self.patient_info.get('检查项目', ''),
            '<检查日期>': self.patient_info.get('检查日期', ''),
            '<分析结果>': str(prediction_result)
        }
        
        report_text = self.report_edit.toPlainText()


        if not report_text.strip():
            QMessageBox.warning(self, '错误', '报告内容为空，无法保存。')
            return

        save_path, _ = QFileDialog.getSaveFileName(self, '保存报告', '', 'Word Files (*.docx)')
        if save_path:
            doc = Document()

            from docx.shared import Pt
            from docx.oxml.ns import qn
            
            # 设置整个文档的默认字体为 宋体
            style = doc.styles['Normal']
            style.font.name = '宋体'
            style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            style.font.size = Pt(12)  # 可选，设置字号

            for line in report_text.split('\n'):
                doc.add_paragraph(line)
            doc.save(save_path)
            QMessageBox.information(self, '成功', '报告已生成并保存。')
        else:
            pass

        # doc = Document(self.template_path)
        # for para in doc.paragraphs:
        #     for key, value in replacements.items():
        #         if key in para.text:
        #             para.text = para.text.replace(key, value)

        # save_path, _ = QFileDialog.getSaveFileName(self, '保存报告', '', 'Word Files (*.docx)')
        # if save_path:
        #     doc.save(save_path)
        #     QMessageBox.information(self, '成功', '报告已生成并保存。')

    def clear_report(self):
        self.report_edit.clear()

    def reset_progress(self):
        self.progress_bar.setValue(0)

    def show_about(self):
        QMessageBox.information(self, '关于', '医学图像处理软件\n版本：1.0\n作者：Chenchao')


if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = MedicalGUI()
    window.show()
    sys.exit(app.exec_())
